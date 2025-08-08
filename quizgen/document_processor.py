from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List


@dataclass
class Document:
    path: Path
    text: str


class DocumentProcessor:
    def __init__(self) -> None:
        pass

    def extract_text(self, file_path: str | Path) -> Document:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {path}")

            suffix = path.suffix.lower()
            if suffix == ".pdf":
                text = self._extract_pdf_text(path)
            elif suffix in (".docx",):
                text = self._extract_docx_text(path)
            elif suffix in (".txt",):
                text = path.read_text(encoding="utf-8", errors="ignore")
            elif suffix in (".html", ".htm"): 
                text = self._extract_html_text(path)
            else:
                # Fallback: try reading as UTF-8 text
                text = path.read_text(encoding="utf-8", errors="ignore")

            text = self._strip_headers_footers(text)
            text = self._normalize_whitespace(text)
            return Document(path=path, text=text)

    def _extract_pdf_text(self, path: Path) -> str:
        try:
            import pdfplumber
        except Exception as exc:
            raise RuntimeError(
                "pdfplumber is required to extract text from PDFs. Install with `pip install pdfplumber`."
            ) from exc

        pages_text: List[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                pages_text.append(page_text)
        return "\n".join(pages_text)

    def _extract_docx_text(self, path: Path) -> str:
        try:
            import docx  # python-docx
        except Exception as exc:
            raise RuntimeError(
                "python-docx is required to extract text from .docx files. Install with `pip install python-docx`."
            ) from exc

        document = docx.Document(str(path))
        paragraphs = [p.text for p in document.paragraphs]
        return "\n".join(paragraphs)

    def _extract_html_text(self, path: Path) -> str:
        try:
            from bs4 import BeautifulSoup
        except Exception as exc:
            raise RuntimeError(
                "beautifulsoup4 is required for HTML parsing. Install with `pip install beautifulsoup4`."
            ) from exc

        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        for script in soup(["script", "style"]):
            script.decompose()
        return soup.get_text(" ")

    def _strip_headers_footers(self, text: str) -> str:
        # Heuristic removal of page numbers / headers / footers
        lines = text.splitlines()
        cleaned: List[str] = []
        for line in lines:
            line_stripped = line.strip()
            if re.fullmatch(r"\d+", line_stripped):
                continue
            if re.match(r"^Page\s+\d+", line_stripped, flags=re.IGNORECASE):
                continue
            cleaned.append(line)
        return "\n".join(cleaned)

    def _normalize_whitespace(self, text: str) -> str:
        text = re.sub(r"\u00A0", " ", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()